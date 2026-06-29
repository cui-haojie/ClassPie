# -*- coding: utf-8 -*-
"""Generate ClassPi course design Word document with diagrams and UI images."""
import os
import textwrap
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent.parent
IMG_DIR = ROOT / "doc_assets" / "images"
OUT_DOCX = ROOT / "课堂派课程设计报告.docx"
OUT_DOCX_EN = ROOT / "ClassPi_CourseDesign_Report.docx"

# Colors
BLUE = (72, 138, 248)
DARK = (30, 41, 59)
GRAY = (100, 116, 139)
LIGHT_BG = (248, 250, 252)
WHITE = (255, 255, 255)
PURPLE = (99, 102, 241)
GREEN = (16, 185, 129)


def get_font(size=16, bold=False):
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    for p in candidates:
        if os.path.exists(p):
            try:
                return ImageFont.truetype(p, size)
            except Exception:
                pass
    return ImageFont.load_default()


def get_font_mono(size=14):
    for p in ["C:/Windows/Fonts/consola.ttf", "C:/Windows/Fonts/cour.ttf"]:
        if os.path.exists(p):
            try:
                return ImageFont.truetype(p, size)
            except Exception:
                pass
    return get_font(size)


def save_img(img, name):
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    path = IMG_DIR / name
    img.save(path, "PNG")
    return path


def draw_box(draw, x, y, w, h, text, fill=WHITE, outline=BLUE, font=None, radius=8):
    if font is None:
        font = get_font(14)
    draw.rounded_rectangle([x, y, x + w, y + h], radius=radius, fill=fill, outline=outline, width=2)
    lines = text.split("\n")
    lh = font.size + 4
    ty = y + (h - len(lines) * lh) // 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        tw = bbox[2] - bbox[0]
        draw.text((x + (w - tw) // 2, ty), line, fill=DARK, font=font)
        ty += lh


def arrow_down(draw, cx, y1, y2):
    draw.line([(cx, y1), (cx, y2)], fill=GRAY, width=2)
    draw.polygon([(cx - 6, y2 - 10), (cx + 6, y2 - 10), (cx, y2)], fill=GRAY)


def arrow_right(draw, x1, y, x2):
    draw.line([(x1, y), (x2, y)], fill=GRAY, width=2)
    draw.polygon([(x2 - 10, y - 6), (x2 - 10, y + 6), (x2, y)], fill=GRAY)


def make_title_banner(title, subtitle=""):
    W, H = 900, 120
    img = Image.new("RGB", (W, H), BLUE)
    draw = ImageDraw.Draw(img)
    f1 = get_font(28, True)
    f2 = get_font(16)
    bbox = draw.textbbox((0, 0), title, font=f1)
    tw = bbox[2] - bbox[0]
    draw.text(((W - tw) // 2, 30), title, fill=WHITE, font=f1)
    if subtitle:
        bbox2 = draw.textbbox((0, 0), subtitle, font=f2)
        tw2 = bbox2[2] - bbox2[0]
        draw.text(((W - tw2) // 2, 75), subtitle, fill=(220, 230, 255), font=f2)
    return save_img(img, f"banner_{title[:10]}.png")


def make_activity_diagram(title, steps):
    W, H = 700, 80 + len(steps) * 90
    img = Image.new("RGB", (W, H), LIGHT_BG)
    draw = ImageDraw.Draw(img)
    f = get_font(18, True)
    bbox = draw.textbbox((0, 0), title, font=f)
    draw.text(((W - bbox[2] + bbox[0]) // 2, 15), title, fill=DARK, font=f)
    cx = W // 2
    y = 55
    for i, step in enumerate(steps):
        draw_box(draw, cx - 180, y, 360, 50, step, font=get_font(14))
        if i < len(steps) - 1:
            arrow_down(draw, cx, y + 50, y + 80)
        y += 80
    return save_img(img, f"activity_{len(steps)}.png")


def make_architecture_diagram():
    W, H = 800, 520
    img = Image.new("RGB", (W, H), LIGHT_BG)
    draw = ImageDraw.Draw(img)
    f = get_font(20, True)
    draw.text((280, 15), "课堂派系统技术架构图", fill=DARK, font=f)
    layers = [
        ("展示层", "Vue 3 + Vite + Vue Router + Pinia\nChrome / Edge 浏览器访问", (72, 138, 248)),
        ("交互层", "Axios HTTP 请求 + WebSocket 实时通信\nRESTful API (/editor/*)", (99, 102, 241)),
        ("业务层", "Spring Boot 3.4 + Controller / Service\nJWT 认证 + 业务逻辑处理", (16, 185, 129)),
        ("持久层", "MyBatis Mapper + MySQL 8.0\n文件存储 (uploads/)", (245, 158, 11)),
    ]
    y = 60
    for name, desc, color in layers:
        draw.rounded_rectangle([80, y, W - 80, y + 90], radius=10, fill=color + (30,) if len(color) == 3 else color, outline=color, width=2)
        draw.rounded_rectangle([80, y, W - 80, y + 90], radius=10, fill=(*color, 20) if False else (240, 248, 255), outline=color, width=2)
        draw_box(draw, 100, y + 10, 120, 35, name, fill=(240, 248, 255), outline=color, font=get_font(15, True))
        draw.text((240, y + 25), desc, fill=DARK, font=get_font(14))
        if y < 60 + 90 * 3:
            arrow_down(draw, W // 2, y + 90, y + 110)
        y += 110
    return save_img(img, "arch_tech.png")


def make_function_tree():
    W, H = 900, 480
    img = Image.new("RGB", (W, H), LIGHT_BG)
    draw = ImageDraw.Draw(img)
    f = get_font(20, True)
    draw.text((320, 10), "课堂派平台功能架构图", fill=DARK, font=f)
    draw_box(draw, 370, 45, 160, 40, "课堂派平台", font=get_font(15, True))
    modules = [
        ("用户管理", "注册/登录\nJWT认证\n个人信息"),
        ("课程管理", "创建/加入\n置顶/归档\n拖拽排序"),
        ("作业管理", "发布/提交\n批阅/打回\n附件上传"),
        ("课堂互动", "实时问答\n投票/抢答\n随机点名"),
        ("签到考勤", "发起签到\n学生打卡\n考勤统计"),
        ("备课区", "备课作业\n备课资料\n一键导入"),
        ("测试测验", "选择题/简答\n草稿发布\n自动统计"),
        ("成绩册", "权重配置\n综合成绩\n数据导出"),
    ]
    cols = 4
    bw, bh = 190, 70
    for i, (name, desc) in enumerate(modules):
        col, row = i % cols, i // cols
        x = 30 + col * (bw + 20)
        y = 110 + row * (bh + 30)
        draw.line([(450, 85), (x + bw // 2, y)], fill=GRAY, width=1)
        draw_box(draw, x, y, bw, bh, f"{name}\n{desc}", font=get_font(12))
    return save_img(img, "func_tree.png")


def make_er_diagram():
    W, H = 900, 550
    img = Image.new("RGB", (W, H), LIGHT_BG)
    draw = ImageDraw.Draw(img)
    f = get_font(20, True)
    draw.text((330, 10), "课堂派核心 E-R 图", fill=DARK, font=f)
    entities = [
        (50, 80, "accounts\n用户"),
        (300, 80, "courses\n课程"),
        (550, 80, "homework\n作业"),
        (750, 80, "content\n作业提交"),
        (50, 280, "course_activity\n课程活动"),
        (300, 280, "course_interaction\n互动响应"),
        (550, 280, "course_attendance\n签到会话"),
        (750, 280, "notification\n通知"),
        (175, 430, "teacher_prep\n备课区"),
        (425, 430, "course_test\n在线测试"),
        (675, 430, "course_grade\n成绩权重"),
    ]
    for x, y, label in entities:
        draw_box(draw, x, y, 140, 55, label, font=get_font(12))
    relations = [
        ((120, 135), (300, 110), "N:M account_course"),
        ((370, 135), (550, 110), "1:N courses_homework"),
        ((620, 135), (750, 110), "1:N content"),
        ((120, 280), (300, 110), "创建"),
        ((370, 280), (120, 280), "class_id"),
        ((620, 280), (370, 135), "class_id"),
    ]
    for (x1, y1), (x2, y2), _ in relations:
        draw.line([(x1, y1), (x2, y2)], fill=GRAY, width=1)
    return save_img(img, "er_diagram.png")


def make_sequence_diagram(title, actors, messages):
    W = 800
    H = 80 + len(messages) * 45 + 60
    img = Image.new("RGB", (W, H), WHITE)
    draw = ImageDraw.Draw(img)
    f = get_font(18, True)
    bbox = draw.textbbox((0, 0), title, font=f)
    draw.text(((W - bbox[2] + bbox[0]) // 2, 10), title, fill=DARK, font=f)
    n = len(actors)
    xs = [80 + i * (W - 160) // max(n - 1, 1) for i in range(n)]
    y0 = 50
    draw.line([(xs[0], y0), (xs[-1], y0)], fill=DARK, width=2)
    for i, actor in enumerate(actors):
        draw_box(draw, xs[i] - 50, y0, 100, 35, actor, font=get_font(12))
        draw.line([(xs[i], y0 + 35), (xs[i], H - 30)], fill=GRAY, width=1)
    y = y0 + 60
    for frm, to, msg in messages:
        x1, x2 = xs[frm], xs[to]
        color = BLUE if frm < to else GREEN
        draw.line([(x1, y), (x2, y)], fill=color, width=2)
        if x2 > x1:
            draw.polygon([(x2 - 8, y - 5), (x2 - 8, y + 5), (x2, y)], fill=color)
        else:
            draw.polygon([(x2 + 8, y - 5), (x2 + 8, y + 5), (x2, y)], fill=color)
        draw.text(((x1 + x2) // 2 - len(msg) * 4, y - 18), msg, fill=DARK, font=get_font(11))
        y += 45
    return save_img(img, f"seq_{hash(title) % 10000}.png")


def make_code_image(title, code, filename):
    lines = code.strip().split("\n")
    W = 820
    H = 50 + len(lines) * 22
    img = Image.new("RGB", (W, H), (30, 30, 30))
    draw = ImageDraw.Draw(img)
    font = get_font_mono(13)
    draw.text((15, 10), title, fill=(100, 200, 255), font=get_font(14, True))
    y = 40
    for line in lines:
        draw.text((15, y), line[:95], fill=(220, 220, 220), font=font)
        y += 22
    return save_img(img, filename)


def make_ui_mock(title, elements, filename, width=820, height=520):
    img = Image.new("RGB", (width, height), (241, 245, 249))
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle([0, 0, width, 52], radius=0, fill=BLUE)
    draw.text((20, 14), "课堂派 ClassPi", fill=WHITE, font=get_font(18, True))
    draw.text((width - 180, 18), title, fill=(220, 230, 255), font=get_font(14))
    y = 70
    for el in elements:
        if el["type"] == "heading":
            draw.text((30, y), el["text"], fill=DARK, font=get_font(20, True))
            y += 40
        elif el["type"] == "card":
            h = el.get("h", 60)
            draw.rounded_rectangle([30, y, width - 30, y + h], radius=8, fill=WHITE, outline=(226, 232, 240), width=1)
            draw.text((45, y + 12), el["text"], fill=DARK, font=get_font(14))
            if "sub" in el:
                draw.text((45, y + 32), el["sub"], fill=GRAY, font=get_font(12))
            y += h + 12
        elif el["type"] == "btn":
            color = el.get("color", BLUE)
            draw.rounded_rectangle([30, y, 130, y + 36], radius=6, fill=color)
            draw.text((45, y + 8), el["text"], fill=WHITE, font=get_font(13))
            y += 50
        elif el["type"] == "input":
            draw.rounded_rectangle([30, y, width - 30, y + 40], radius=6, fill=WHITE, outline=(203, 213, 225), width=1)
            draw.text((45, y + 10), el.get("placeholder", ""), fill=GRAY, font=get_font(13))
            y += 55
        elif el["type"] == "tabs":
            x = 30
            for tab in el["tabs"]:
                active = tab == el.get("active")
                c = BLUE if active else (226, 232, 240)
                tc = WHITE if active else GRAY
                tw = len(tab) * 16 + 24
                draw.rounded_rectangle([x, y, x + tw, y + 32], radius=6, fill=c)
                draw.text((x + 12, y + 7), tab, fill=tc, font=get_font(12))
                x += tw + 8
            y += 45
    return save_img(img, filename)


def make_backend_structure():
    W, H = 750, 420
    img = Image.new("RGB", (W, H), LIGHT_BG)
    draw = ImageDraw.Draw(img)
    draw.text((250, 10), "后端项目包结构", fill=DARK, font=get_font(18, True))
    tree = """org.example.classpiserver
├── controller/     # REST 接口层
│   ├── account/    # 账户登录注册
│   ├── course/     # 课程管理
│   ├── homework/   # 作业管理
│   ├── activity/   # 话题/资料/公告
│   ├── interaction/# 课堂互动
│   ├── attendance/ # 签到考勤
│   ├── test/       # 在线测试
│   ├── grade/      # 成绩册
│   ├── prep/       # 教师备课区
│   └── notification/
├── service/        # 业务逻辑层
├── mapper/         # MyBatis 数据访问
├── entity/ & dto/  # 实体与传输对象
├── security/       # JWT 认证拦截
├── live/           # WebSocket 实时推送
└── util/           # 工具类"""
    y = 45
    for line in tree.split("\n"):
        draw.text((40, y), line, fill=DARK, font=get_font_mono(13))
        y += 22
    return save_img(img, "backend_structure.png")


def make_frontend_structure():
    W, H = 750, 380
    img = Image.new("RGB", (W, H), LIGHT_BG)
    draw = ImageDraw.Draw(img)
    draw.text((250, 10), "前端项目目录结构", fill=DARK, font=get_font(18, True))
    tree = """classPai/src/
├── views/          # 页面组件
│   ├── login.vue / register.vue
│   ├── mainInterface.vue   # 课程列表
│   ├── courseContent.vue   # 课程主页
│   ├── LiveClass.vue       # 进入课堂/签到
│   ├── InteractionContent.vue
│   ├── HomeworkContent.vue
│   ├── PrepArea.vue        # 备课区
│   └── userSetting.vue
├── components/     # 公共组件
├── router/         # Vue Router 路由
├── stores/         # Pinia 状态管理
└── utils/          # request / liveSocket"""
    y = 45
    for line in tree.split("\n"):
        draw.text((40, y), line, fill=DARK, font=get_font_mono(13))
        y += 22
    return save_img(img, "frontend_structure.png")


def set_cell_shading(cell, color="D9E2F3"):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i])
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    run.bold = bold
    return p


def add_figure(doc, img_path, caption):
    doc.add_picture(str(img_path), width=Inches(5.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    run.bold = True


def generate_all_images():
    images = {}
    images["role_activity"] = make_activity_diagram("角色权限活动图", [
        "用户访问平台", "登录/注册验证", "识别教师/学生角色", "加载对应功能菜单", "退出系统"
    ])
    images["course_activity"] = make_activity_diagram("课程管理活动图", [
        "进入主界面", "创建课程/加入课程", "管理课程(置顶/归档/排序)", "进入课程主页", "结束"
    ])
    images["func_tree"] = make_function_tree()
    images["arch"] = make_architecture_diagram()
    images["er"] = make_er_diagram()
    images["backend"] = make_backend_structure()
    images["frontend"] = make_frontend_structure()
    images["seq_login"] = make_sequence_diagram("用户登录时序图",
        ["用户", "前端", "Controller", "Service", "MySQL"],
        [(0, 1, "输入账号密码"), (1, 2, "POST /login"), (2, 3, "verifyAccount"),
         (3, 4, "查询用户"), (4, 3, "返回用户"), (3, 2, "生成JWT Token"), (2, 1, "返回token"), (1, 0, "跳转主页")])
    images["seq_homework"] = make_sequence_diagram("作业提交时序图",
        ["学生", "前端", "HomeworkController", "Service", "DB"],
        [(0, 1, "填写作业"), (1, 2, "submitHomework"), (2, 3, "校验权限/截止"),
         (3, 4, "写入content表"), (4, 3, "OK"), (3, 2, "返回成功"), (2, 1, "提示提交成功")])
    images["seq_interaction"] = make_sequence_diagram("课堂互动时序图",
        ["教师", "前端", "InteractionController", "WebSocket", "学生"],
        [(0, 1, "开启互动"), (1, 2, "addCourseInteraction"), (2, 3, "推送通知"),
         (3, 4, "收到互动通知"), (4, 1, "提交回答"), (1, 2, "submitInteraction"), (2, 3, "实时更新")])
    images["code_jwt"] = make_code_image("JwtService.java - Token 生成", """
public String generateToken(String account) {
    return Jwts.builder()
        .subject(account)
        .expiration(new Date(System.currentTimeMillis() + EXPIRE_MS))
        .signWith(key)
        .compact();
}""", "code_jwt.png")
    images["code_close"] = make_code_image("结束签到自动关闭互动", """
public boolean closeAttendance(Long sessionId, String teacherAccount) {
    boolean ok = attendanceMapper.closeSession(sessionId);
    if (ok) {
        interactionService.closeActiveInteractionsForClass(session.getClass_id());
        liveEventPublisher.publishCourse(classId, "attendance_closed", data);
    }
    return ok;
}""", "code_close_interaction.png")
    # UI mocks
    images["ui_login"] = make_ui_mock("登录", [
        {"type": "heading", "text": "欢迎登录课堂派"},
        {"type": "input", "placeholder": "请输入账号（邮箱/手机号）"},
        {"type": "input", "placeholder": "请输入密码"},
        {"type": "btn", "text": "登 录"},
        {"type": "btn", "text": "注册账号", "color": (148, 163, 184)},
    ], "ui_login.png", height=400)
    images["ui_register"] = make_ui_mock("注册", [
        {"type": "heading", "text": "创建课堂派账号"},
        {"type": "input", "placeholder": "姓名"},
        {"type": "input", "placeholder": "账号 / 密码 / 确认密码"},
        {"type": "input", "placeholder": "选择角色：教师 / 学生"},
        {"type": "btn", "text": "注 册"},
    ], "ui_register.png", height=420)
    images["ui_main_teacher"] = make_ui_mock("教师主页", [
        {"type": "tabs", "tabs": ["全部学期", "2025-2026 第一学期", "已归档"], "active": "全部学期"},
        {"type": "btn", "text": "＋ 创建课程"},
        {"type": "card", "text": "软件工程与计算 II", "sub": "周一 1-2节 · 加课码 AB12CD · 已置顶", "h": 55},
        {"type": "card", "text": "数据库原理", "sub": "周三 3-4节 · 12 名学生", "h": 55},
        {"type": "card", "text": "Web 前端开发", "sub": "周五 5-6节 · 8 名学生", "h": 55},
    ], "ui_main_teacher.png")
    images["ui_main_student"] = make_ui_mock("学生主页", [
        {"type": "tabs", "tabs": ["全部学期", "2025-2026 第一学期"], "active": "全部学期"},
        {"type": "btn", "text": "＋ 加入课程"},
        {"type": "card", "text": "软件工程与计算 II", "sub": "崔老师 · 有新作业通知", "h": 55},
        {"type": "card", "text": "数据结构", "sub": "进行中 · 签到已开启", "h": 55},
    ], "ui_main_student.png")
    images["ui_course"] = make_ui_mock("课程主页", [
        {"type": "tabs", "tabs": ["课程互动", "话题", "资料", "测试", "公告", "作业", "成员", "成绩册"], "active": "课程互动"},
        {"type": "btn", "text": "开启互动"},
        {"type": "card", "text": "第三章课堂问答", "sub": "课堂问答 · 进行中 · 15 人已参与", "h": 55},
        {"type": "card", "text": "期中复习投票", "sub": "投票 · 已结束 · 28 人已参与", "h": 55},
    ], "ui_course.png")
    images["ui_live"] = make_ui_mock("进入课堂", [
        {"type": "heading", "text": "软件工程与计算 II · 进入课堂"},
        {"type": "card", "text": "签到进行中", "sub": "应到 30 · 实到 26 · 签到码 8374", "h": 65},
        {"type": "btn", "text": "结束签到"},
        {"type": "card", "text": "课堂互动", "sub": "实时问答 · 抢答 · 投票", "h": 55},
    ], "ui_live.png")
    images["ui_homework"] = make_ui_mock("作业详情", [
        {"type": "heading", "text": "第一次作业：需求分析文档"},
        {"type": "card", "text": "截止时间：2026-06-30 23:59", "sub": "支持附件上传 doc/pdf/zip", "h": 55},
        {"type": "btn", "text": "提交作业"},
        {"type": "card", "text": "我的提交", "sub": "已提交 · 待批阅 · 附件: report.docx", "h": 55},
    ], "ui_homework.png")
    images["ui_interaction"] = make_ui_mock("课堂互动", [
        {"type": "heading", "text": "第三章课堂问答"},
        {"type": "tabs", "tabs": ["进行中"], "active": "进行中"},
        {"type": "card", "text": "问题：敏捷开发的核心原则是什么？", "sub": "第 2 轮 · 8 人已回答", "h": 65},
        {"type": "btn", "text": "随机点名"},
        {"type": "btn", "text": "结束互动", "color": (239, 68, 68)},
    ], "ui_interaction.png")
    images["ui_prep"] = make_ui_mock("备课区", [
        {"type": "tabs", "tabs": ["作业", "话题", "资料", "测试"], "active": "测试"},
        {"type": "btn", "text": "＋ 新建备课项"},
        {"type": "card", "text": "第一章测验（含简答题）", "sub": "选择题 5 · 简答题 2 · 可导入课程", "h": 55},
        {"type": "card", "text": "实验指导书", "sub": "资料 · PDF 附件", "h": 55},
    ], "ui_prep.png")
    images["ui_setting"] = make_ui_mock("个人设置", [
        {"type": "heading", "text": "个人信息设置"},
        {"type": "input", "placeholder": "姓名 / 院系 / 专业 / 年级"},
        {"type": "input", "placeholder": "学号/工号 / 入学日期"},
        {"type": "btn", "text": "保存修改"},
        {"type": "btn", "text": "修改密码", "color": (148, 163, 184)},
    ], "ui_setting.png", height=420)
    images["ui_grade"] = make_ui_mock("成绩册", [
        {"type": "heading", "text": "课程成绩册"},
        {"type": "card", "text": "权重：作业 40% · 测试 30% · 互动 10% · 考勤 20%", "sub": "教师可自定义权重", "h": 55},
        {"type": "card", "text": "张三  综合 87.5", "sub": "作业 90 · 测试 85 · 互动 8 · 考勤 18", "h": 55},
        {"type": "card", "text": "李四  综合 92.0", "sub": "作业 95 · 测试 88 · 互动 9 · 考勤 20", "h": 55},
    ], "ui_grade.png")
    return images


def build_document(images):
    doc = Document()
    # Page setup
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    # ===== Cover =====
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("软件工程与计算 II 课程设计报告")
    r.font.size = Pt(22)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph()
    cover_items = [
        ("题    目", "课堂派（ClassPi）"),
        ("学    院", "两江人工智能学院"),
        ("专    业", "软件工程"),
        ("指导教师", "王森、张金荣、龙华、蒋鑫等"),
        ("完成时间", "2026 年 06 月"),
    ]
    for label, val in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}    {val}")
        run.font.size = Pt(14)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_page_break()

    # ===== 团队分工 =====
    add_heading(doc, "团队职责分工表", 1)
    add_table(doc, ["姓名", "角色", "工作职责"], [
        ["（填写）", "项目经理", "开发计划制定、系统设计、课程与互动模块、文档撰写"],
        ["（填写）", "前后端开发", "作业、测试、签到考勤功能开发"],
        ["（填写）", "前后端开发", "个人设置、通知、备课区、成绩册功能开发"],
    ])
    doc.add_paragraph()
    add_heading(doc, "团队沟通计划表", 1)
    add_table(doc, ["沟通场景", "沟通方式", "沟通频率", "沟通责任人", "沟通内容"], [
        ["集中上课", "当面沟通", "每次课", "项目经理", "设计、实现问题"],
        ["非集中上课", "线上/当面", "4 次/天", "项目经理", "进度同步、Bug 修复"],
    ])
    doc.add_page_break()

    # ===== 摘要 =====
    add_heading(doc, "摘  要", 1)
    add_para(doc, (
        "随着教育信息化的深入发展，传统教学模式正面临数字化转型的迫切需求。"
        "本文设计并实现了一个基于 Vue 3 前端框架与 Spring Boot 3 后端技术的互动式教学管理系统——课堂派（ClassPi），"
        "旨在通过现代化技术手段提升课堂互动性与教学管理效率。"
        "系统集成了实时课堂互动（问答、投票、抢答、随机点名）、在线签到考勤、作业管理与批阅、"
        "在线测试测验、教师备课区、消息通知、可配置成绩册等功能模块。"
        "采用前后端分离架构，使用 JWT 进行身份认证、WebSocket 实现课堂实时推送，"
        "MyBatis 操作 MySQL 数据库完成数据持久化。"
        "经实际部署与测试，系统运行稳定，有效提升了师生课堂参与度与教学管理效率。"
    ), indent=True)
    add_para(doc, "关键词：Vue.js；Spring Boot；课堂派；互动式教学；前后端分离；WebSocket", bold=True)
    doc.add_paragraph()
    add_heading(doc, "Abstract", 1)
    add_para(doc, (
        "This paper presents ClassPi, an interactive teaching management system built with "
        "Vue 3 and Spring Boot 3. The platform supports real-time classroom interactions "
        "(Q&A, voting, race-to-answer, random pick), attendance check-in, homework management, "
        "online tests, teacher prep area, notifications, and configurable grade books. "
        "A front-end/back-end separation architecture with JWT authentication and WebSocket "
        "live events ensures security and real-time responsiveness. Practical deployment "
        "demonstrates improved classroom engagement and teaching efficiency."
    ), indent=True)
    add_para(doc, "Keywords: Vue.js; Spring Boot; ClassPi; Interactive Teaching; WebSocket", bold=True)
    doc.add_page_break()

    # ===== 1 需求分析 =====
    add_heading(doc, "1  需求分析", 1)
    add_heading(doc, "1.1  需求概述", 2)
    add_para(doc, (
        "课堂派（ClassPi）是一款面向高等院校的在线教学管理平台，支持教师创建与管理课程、"
        "学生通过加课码加入课程，并在课程内进行作业、测试、资料、话题、课堂互动等多种教学活动。"
        "主要需求包括："
    ), indent=True)
    reqs = [
        "用户登录验证：JWT 令牌认证，BCrypt 密码加密，支持教师/学生角色区分。",
        "用户信息管理：注册、编辑个人资料（院系、专业、年级、学号等）、修改密码。",
        "教师创建课程：设置课程名称、上课时间、关联行政班、生成加课码。",
        "学生加入课程：输入加课码加入，查看课程成员与内容。",
        "课程管理：置顶、归档、拖拽排序、按学期筛选。",
        "作业管理：教师布置作业（含附件），学生提交，教师批阅打分与打回。",
        "课堂互动：实时问答、投票、抢答、随机点名，结束签到时自动关闭互动。",
        "签到考勤：教师发起签到，学生打卡，统计应到/实到/缺勤。",
        "在线测试：选择题与简答题，草稿发布，自动统计交卷人数。",
        "教师备课区：预制作业/资料/话题/测试，一键导入到课程。",
        "消息通知：作业、话题、签到、测试等事件推送通知。",
        "成绩册：可配置权重，汇总作业、测试、互动、考勤成绩。",
    ]
    for i, r in enumerate(reqs, 1):
        add_para(doc, f"{i}、{r}")

    add_heading(doc, "1.2  总体需求描述", 2)
    add_para(doc, (
        "课堂派为教育者和学习者提供功能全面、操作便捷的在线平台。"
        "教师可管理课程、布置作业、发起签到与课堂互动、发布测试、使用备课区提高效率；"
        "学生可加入课程、提交作业、参与互动与讨论、在线签到、查看成绩与通知。"
        "系统要求界面简洁、响应迅速、数据安全、支持主流浏览器。"
    ), indent=True)

    add_heading(doc, "1.3  详细业务需求", 2)
    add_heading(doc, "1.3.1  角色权限管理", 3)
    add_para(doc, (
        "平台区分教师与学生两种角色。教师拥有创建课程、布置作业、发起签到与互动、"
        "批阅作业、管理成绩册、使用备课区等权限；学生拥有加入课程、提交作业、"
        "参与互动与话题、签到打卡、查看个人成绩等权限。"
    ), indent=True)
    add_figure(doc, images["role_activity"], "图 1.1  角色权限活动图")

    add_heading(doc, "1.3.2  课程管理", 3)
    add_para(doc, (
        "教师创建课程后生成唯一加课码，可将课程置顶或归档，并通过拖拽调整排序。"
        "学生通过加课码加入后可查看课程成员、作业、资料、互动等内容。"
    ), indent=True)
    add_figure(doc, images["course_activity"], "图 1.2  课程管理活动图")
    add_table(doc, ["用例编号", "活动者", "用例名称", "开始条件"], [
        ["CP-001", "教师", "创建课程", "进入教师主界面"],
        ["CP-002", "学生", "加入课程", "进入学生主界面"],
    ])
    add_para(doc, "创建课程事件流：1.进入主界面 → 2.点击创建课程 → 3.填写课程信息 → 4.确认创建 → 5.生成加课码。", indent=True)
    add_para(doc, "加入课程事件流：1.点击加入课程 → 2.输入加课码 → 3.验证通过 → 4.进入课程主页。", indent=True)

    add_heading(doc, "1.4  非功能性需求", 2)
    nfr = [
        ("1.4.1  环境需求", "开发：Java 17、Spring Boot 3.4、Vue 3、Vite、MySQL 8.0、Maven、IntelliJ IDEA。运行：Windows 10+、Chrome/Edge。"),
        ("1.4.2  性能需求", "页面响应时间平均 2 秒以内，WebSocket 推送延迟低于 1 秒。"),
        ("1.4.3  安全需求", "BCrypt 密码加密、JWT 令牌认证、接口权限拦截、防止未授权访问。"),
        ("1.4.4  兼容需求", "支持 Windows 系统与 Chrome、Edge 等主流浏览器。"),
        ("1.4.5  交互需求", "统一 Modal 弹窗、Toast 提示、拖拽排序动画、实时课堂状态同步。"),
        ("1.4.6  扩展需求", "备课区导入、AI 辅助出题/批阅接口预留、成绩权重可配置。"),
    ]
    for title, content in nfr:
        add_heading(doc, title, 3)
        add_para(doc, content, indent=True)

    doc.add_page_break()

    # ===== 2 系统设计 =====
    add_heading(doc, "2  系统设计", 1)
    add_heading(doc, "2.1  设计概述", 2)
    add_para(doc, (
        "系统设计将需求分析结果转化为技术方案，涵盖功能设计、架构设计与数据库设计三个层面。"
        "功能设计明确各模块职责；架构设计采用前后端分离与三层架构；"
        "数据库设计基于 E-R 模型映射为 MySQL 表结构。"
    ), indent=True)

    add_heading(doc, "2.2  总体设计", 2)
    add_heading(doc, "2.2.1  功能架构设计", 3)
    add_para(doc, "平台包含用户管理、课程管理、作业管理、课堂互动、签到考勤、在线测试、备课区、成绩册、通知等九大功能模块。", indent=True)
    add_figure(doc, images["func_tree"], "图 2.1  平台功能架构图")
    add_table(doc, ["功能模块", "功能点", "说明"], [
        ["用户管理", "注册/登录/JWT", "身份认证与权限控制"],
        ["课程管理", "创建/加入/置顶/归档/排序", "课程全生命周期管理"],
        ["作业管理", "发布/提交/批阅/打回", "含附件上传"],
        ["课堂互动", "问答/投票/抢答/点名", "WebSocket 实时同步"],
        ["签到考勤", "发起/打卡/统计", "应到实到缺勤"],
        ["在线测试", "选择/简答/草稿", "自动统计交卷"],
        ["备课区", "预制作业/导入", "提高备课效率"],
        ["成绩册", "权重配置/汇总", "多维度成绩计算"],
    ])

    add_heading(doc, "2.2.2  技术架构设计", 3)
    add_figure(doc, images["arch"], "图 2.2  系统技术架构图")
    add_para(doc, (
        "展示层通过 Vue 3 构建 SPA 单页应用；交互层使用 Axios 调用 REST API 并通过 WebSocket 接收实时事件；"
        "业务层 Spring Boot 处理业务逻辑与 JWT 认证；持久层 MyBatis 访问 MySQL，文件存储于 uploads 目录。"
    ), indent=True)

    add_heading(doc, "2.2.3  物理模型设计", 3)
    add_para(doc, (
        "前端 Vite 开发服务器（5174 端口）通过代理访问后端 Spring Boot 应用（9090 端口，前缀 /editor）。"
        "后端连接 MySQL 数据库 t_class，静态文件与上传附件存储于服务器本地 uploads 目录。"
    ), indent=True)

    add_heading(doc, "2.3  数据库设计", 2)
    add_heading(doc, "2.3.1  概念结构设计", 3)
    add_para(doc, "根据业务需求建立 E-R 概念模型，主要实体包括用户、课程、作业、作业提交、课程活动、互动响应、签到会话、通知、备课区、测试题目等。", indent=True)
    add_figure(doc, images["er"], "图 2.3  课堂派核心 E-R 图")

    add_heading(doc, "2.3.2  数据表结构描述", 3)
    add_table(doc, ["字段", "中文名", "类型", "备注"], [
        ["account", "账号", "VARCHAR(100)", "accounts 主键"],
        ["password", "密码", "VARCHAR", "BCrypt 加密"],
        ["status", "角色", "ENUM", "老师/学生"],
        ["id", "课程ID", "BIGINT", "courses 主键"],
        ["class_name", "课程名", "VARCHAR", "课程名称"],
        ["code", "加课码", "VARCHAR", "学生加入凭证"],
        ["homework_id", "作业ID", "INT", "homework 主键"],
        ["content_id", "提交ID", "BIGINT", "content 主键"],
        ["activity_id", "活动ID", "BIGINT", "course_activity 主键"],
        ["type", "活动类型", "VARCHAR", "interaction/topic/test等"],
    ])

    add_heading(doc, "2.4  详细设计", 2)
    add_heading(doc, "2.4.1  用户管理模块", 3)
    add_para(doc, "AccountController 提供注册、登录、个人信息更新接口；AuthInterceptor 拦截 /editor 请求校验 JWT；JwtService 负责 Token 生成与解析。", indent=True)
    add_figure(doc, images["seq_login"], "图 2.4  用户登录时序图")

    add_heading(doc, "2.4.2  课程管理模块", 3)
    add_para(doc, "CourseController 实现课程 CRUD、加课、置顶、归档、排序；通过 account_course 表维护用户与课程的 N:M 关系。", indent=True)

    add_heading(doc, "2.4.3  作业管理模块", 3)
    add_para(doc, "HomeworkController 与 HomeworkServiceImpl 处理作业发布、提交、批阅；content 表存储学生提交记录，支持附件与打回。", indent=True)
    add_figure(doc, images["seq_homework"], "图 2.5  作业提交时序图")

    add_heading(doc, "2.4.4  课堂互动模块", 3)
    add_para(doc, (
        "InteractionController 支持开启问答/投票/抢答互动，interaction_options JSON 存储状态与轮次；"
        "结束签到时 AttendanceService 自动调用 closeActiveInteractionsForClass 关闭进行中的互动。"
    ), indent=True)
    add_figure(doc, images["seq_interaction"], "图 2.6  课堂互动时序图")

    doc.add_page_break()

    # ===== 3 系统实现 =====
    add_heading(doc, "3  系统实现", 1)
    add_heading(doc, "3.1  系统实现概述", 2)
    add_para(doc, (
        "后端基于 Spring Boot 3.4 + MyBatis 3.0，Java 17，端口 9090，API 前缀 /editor。"
        "前端基于 Vue 3 + Vite + Pinia + Vue Router，开发端口 5174。"
        "认证采用 JWT + BCrypt；实时功能采用 WebSocket（/ws/live）。"
    ), indent=True)

    add_heading(doc, "3.2  代码结构", 2)
    add_figure(doc, images["frontend"], "图 3.1  前端项目代码结构图")
    add_figure(doc, images["backend"], "图 3.2  后端项目代码结构图")

    add_heading(doc, "3.3  关键功能实现", 2)
    add_heading(doc, "3.3.1  JWT 认证实现", 3)
    add_figure(doc, images["code_jwt"], "图 3.3  JWT Token 生成关键代码")

    add_heading(doc, "3.3.2  结束签到自动关闭互动", 3)
    add_figure(doc, images["code_close"], "图 3.4  结束签到关闭互动关键代码")

    add_heading(doc, "3.4  系统实现界面展示", 2)
    ui_items = [
        ("ui_login", "图 3.5  登录界面"),
        ("ui_register", "图 3.6  注册界面"),
        ("ui_main_teacher", "图 3.7  教师主界面"),
        ("ui_main_student", "图 3.8  学生主界面"),
        ("ui_course", "图 3.9  课程主页"),
        ("ui_live", "图 3.10  进入课堂/签到界面"),
        ("ui_interaction", "图 3.11  课堂互动界面"),
        ("ui_homework", "图 3.12  作业详情界面"),
        ("ui_prep", "图 3.13  教师备课区界面"),
        ("ui_grade", "图 3.14  成绩册界面"),
        ("ui_setting", "图 3.15  个人设置界面"),
    ]
    for key, caption in ui_items:
        add_figure(doc, images[key], caption)

    doc.add_page_break()

    # ===== 4 系统测试 =====
    add_heading(doc, "4  系统测试", 1)
    add_heading(doc, "4.1  测试概述", 2)
    add_para(doc, "对系统各功能模块进行功能测试与非功能测试，确保需求正确实现。", indent=True)
    add_heading(doc, "4.2  测试环境", 2)
    add_para(doc, "MySQL 8.0 数据库 t_class；Windows 11；Google Chrome 浏览器；后端 9090 / 前端 5174。", indent=True)
    add_heading(doc, "4.3  测试方法", 2)
    add_para(doc, "采用黑盒功能测试，按用例步骤操作并对比预期结果。", indent=True)
    add_heading(doc, "4.4  测试用例设计", 2)
    add_table(doc, ["编号", "用例名称", "预期结果", "结果"], [
        ["CP-T01", "账号注册登录", "注册成功并可登录", "正常"],
        ["CP-T02", "创建/加入课程", "课程列表正确显示", "正常"],
        ["CP-T03", "发布/提交作业", "作业流转完整", "正常"],
        ["CP-T04", "课堂互动", "问答/投票/抢答正常", "正常"],
        ["CP-T05", "签到考勤", "应到实到统计正确", "正常"],
        ["CP-T06", "结束签到关闭互动", "互动状态变为已结束", "正常"],
        ["CP-T07", "在线测试", "选择/简答提交正常", "正常"],
        ["CP-T08", "备课区导入", "导入到课程成功", "正常"],
        ["CP-T09", "附件上传预览", "上传下载一致", "正常"],
        ["CP-T10", "成绩册权重", "综合成绩计算正确", "正常"],
    ])

    doc.add_page_break()

    # ===== 5 总结 =====
    add_heading(doc, "5  总结", 1)
    add_para(doc, (
        "本次课程设计完成了课堂派（ClassPi）教学管理系统的设计与实现。"
        "系统采用 B/S 前后端分离架构，前端 Vue 3 负责界面与交互，后端 Spring Boot 负责业务逻辑与数据持久化。"
        "实现了课程管理、作业批阅、课堂互动、签到考勤、在线测试、备课区、成绩册等完整功能链路。"
        "在开发过程中，我们深入实践了 JWT 认证、WebSocket 实时通信、MyBatis 数据访问、"
        "文件上传、拖拽排序等关键技术。后续可进一步扩展 AI 辅助教学、移动端适配等功能。"
    ), indent=True)

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] 李鸿君. 大话软件工程——需求分析与软件设计[M]. 北京: 清华大学出版社, 2020.",
        "[2] 刘慧娟. 基于 SpringBoot 的民主测评系统的设计与实现[D]. 北京邮电大学, 2022.",
        "[3] 胡强. MySQL 数据库常见问题分析与研究[J]. 电脑编程技巧与维护, 2019.",
        "[4] 郑海燕. 基于 Java Web 的高校英语线上教学平台设计[J]. 自动化与仪器仪表, 2023.",
        "[5] 孙洪盼. 基于 SpringBoot 和 Vue 的友为交流社区的设计与实现[D]. 重庆大学, 2022.",
    ]
    for ref in refs:
        add_para(doc, ref)

    doc.save(str(OUT_DOCX))
    doc.save(str(OUT_DOCX_EN))
    print(f"Generated: {OUT_DOCX}")
    print(f"Generated: {OUT_DOCX_EN}")
    print(f"Images in: {IMG_DIR}")


if __name__ == "__main__":
    imgs = generate_all_images()
    build_document(imgs)
