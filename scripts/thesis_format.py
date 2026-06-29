# -*- coding: utf-8 -*-
"""Chinese graduation thesis / course design report formatting helpers."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def setup_document_styles(doc):
    """Apply thesis-style Normal and Heading styles."""
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0.74)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for i, (size, bold, indent, align, sb, sa) in enumerate([
        (16, True, 0, WD_ALIGN_PARAGRAPH.CENTER, Pt(12), Pt(12)),   # H1 章
        (14, True, 0, WD_ALIGN_PARAGRAPH.LEFT, Pt(6), Pt(6)),      # H2 节
        (12, True, 0, WD_ALIGN_PARAGRAPH.LEFT, Pt(3), Pt(3)),        # H3 小节
    ], start=1):
        style = doc.styles[f"Heading {i}"]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        spf = style.paragraph_format
        spf.first_line_indent = Cm(indent)
        spf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        spf.space_before = sb
        spf.space_after = sa
        spf.alignment = align
        spf.keep_with_next = True


def setup_page(section, header_text=None, footer_page_num=True):
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    if header_text:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(header_text)
        set_run_font(run, "宋体", 9)
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)
    if footer_page_num:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number_field(fp)


def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)
    run3 = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)
    run4 = paragraph.add_run("1")
    set_run_font(run4, "宋体", 10.5)
    run5 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def add_toc_field(doc):
    """Insert Word TOC field — user opens doc and updates field (F9) for page numbers."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    run2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    run2._r.append(instr)
    run3 = p.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)
    run4 = p.add_run("（请在 Word 中右键目录 → 更新域，以生成页码）")
    set_run_font(run4, "宋体", 10.5)
    run5 = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def add_chapter(doc, title, page_break=True):
    if page_break:
        doc.add_page_break()
    h = doc.add_heading(title, level=1)
    for r in h.runs:
        set_run_font(r, "黑体", 16, True)


def add_section(doc, title):
    h = doc.add_heading(title, level=2)
    for r in h.runs:
        set_run_font(r, "黑体", 14, True)


def add_subsection(doc, title):
    h = doc.add_heading(title, level=3)
    for r in h.runs:
        set_run_font(r, "黑体", 12, True)


def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)
    return p


def add_body_no_indent(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.alignment = align
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)
    return p


def add_numbered_items(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(f"{i}、{item}")
        set_run_font(run, "宋体", 12)


def add_figure(doc, img_path, caption, width=Inches(5.2)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(img_path), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    set_run_font(r, "宋体", 10.5, bold=True)


def set_cell_shading(cell, color="D9E2F3"):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_caption(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(caption)
    set_run_font(r, "宋体", 10.5, bold=True)


def add_table(doc, headers, rows, caption=None):
    if caption:
        add_table_caption(doc, caption)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i])
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            for r in p.runs:
                set_run_font(r, "宋体", 10.5, True)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Cm(0)
                for r in p.runs:
                    set_run_font(r, "宋体", 10.5)
    doc.add_paragraph()
    return table


def add_use_case_table(doc, number, actor, name, precondition, steps, expected, caption):
    add_table(doc,
              ["用例编号", "版本号", "用例名称", "活动者"],
              [[number, "V1.0", name, actor]],
              caption=caption)
    add_table(doc, ["开始条件", "事件流", "结束条件", "预期输出"],
              [[precondition, steps, "退出系统", expected]])


def add_test_case_table(doc, number, name, steps, input_data, expected, caption):
    add_table(doc,
              ["测试用例编号", "版本号", "用例名称", "测试环境"],
              [[number, "V1.0", name, "Windows 11、Chrome、MySQL 8.0"]],
              caption=caption)
    add_table(doc, ["测试步骤", "输入数据", "预期输出", "测试结果"],
              [[steps, input_data, expected, "正常"]])
