# -*- coding: utf-8 -*-
"""
Generate ClassPi graduation-thesis-format course design report (~60+ pages).
Run: py -3.12 scripts/generate_thesis_doc.py
"""
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(Path(__file__).parent))

from thesis_format import (
    setup_document_styles, setup_page, add_toc_field,
    add_chapter, add_section, add_subsection, add_body, add_body_no_indent,
    add_numbered_items, add_figure, add_table, add_use_case_table, add_test_case_table,
    set_run_font,
)
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from thesis_content import (
    module_design_user, module_design_course,
    module_design_homework, module_design_interaction,
    design_overview_paragraphs, impl_overview_paragraphs,
    test_overview_paragraphs, summary_paragraphs,
)
from generate_course_doc import generate_all_images
import subprocess

OUT = ROOT / "ClassPi_Thesis_Report_v2.docx"
OUT_ALT = ROOT / "ClassPi_Thesis_Report.docx"
OUT_CN = ROOT / "课堂派课程设计报告.docx"
HEADER = "两江人工智能学院  软件工程与计算 II 课程设计报告"
SCREENSHOT_DIR = ROOT / "doc_assets" / "screenshots"


def capture_project_assets():
    """Capture real UI + code screenshots from running ClassPi."""
    scripts_dir = Path(__file__).parent
    for name in ("capture_screenshots.py", "capture_code_images.py"):
        script = scripts_dir / name
        print(f"Running {name}...")
        subprocess.run([sys.executable, str(script)], check=False)


def load_project_images():
    """Merge architecture diagrams with real project screenshots."""
    diagrams = generate_all_images()
    images = dict(diagrams)

    shot_map = {
        "ui_login": "ui_login.png",
        "ui_register": "ui_register.png",
        "ui_main_teacher": "ui_main_teacher.png",
        "ui_main_student": "ui_main_student.png",
        "ui_course": "ui_course.png",
        "ui_course_student": "ui_course_student.png",
        "ui_live": "ui_live.png",
        "ui_interaction": "ui_interaction.png",
        "ui_homework": "ui_homework.png",
        "ui_prep": "ui_prep.png",
        "ui_grade": "ui_grade.png",
        "ui_setting": "ui_setting.png",
    }
    for key, fname in shot_map.items():
        p = SCREENSHOT_DIR / fname
        if p.exists():
            images[key] = p

    code_map = {
        "code_jwt": "code_jwt.png",
        "code_close": "code_close_interaction.png",
        "code_interaction": "code_interaction.png",
        "code_course": "code_course.png",
        "code_websocket": "code_websocket.png",
        "code_interaction_vue": "code_interaction_vue.png",
    }
    for key, fname in code_map.items():
        p = SCREENSHOT_DIR / fname
        if p.exists():
            images[key] = p

    return images


def add_lines(doc, lines):
    for line in lines:
        if line.startswith("（") or line.startswith("功能") or line.startswith("静态") or line.startswith("动态"):
            add_body_no_indent(doc, line)
        else:
            add_body(doc, line)


def ref_img(name, fallback_key, images):
    p = REF_MEDIA / name
    if p.exists():
        return p
    return images.get(fallback_key)


DB_TABLES = [
    ("accounts", "用户表", [
        ("account", "账号", "VARCHAR(100)", "主键", "登录账号"),
        ("name", "姓名", "VARCHAR(100)", "", "显示名称"),
        ("status", "角色", "ENUM", "", "老师/学生"),
        ("password", "密码", "VARCHAR", "", "BCrypt 加密"),
        ("mechanism", "机构", "VARCHAR(50)", "", "学校/机构名"),
        ("status_number", "学号/工号", "VARCHAR(255)", "", ""),
        ("avatar_url", "头像", "VARCHAR(512)", "", "头像路径"),
        ("department", "院系", "VARCHAR", "", "可编辑"),
        ("major", "专业", "VARCHAR", "", "可编辑"),
    ]),
    ("courses", "课程表", [
        ("id", "课程ID", "BIGINT", "主键", "自增"),
        ("teacher_account", "教师账号", "VARCHAR(100)", "外键", ""),
        ("class_name", "课程名称", "VARCHAR", "", ""),
        ("class_time", "上课时间", "VARCHAR(512)", "", "可多段"),
        ("code", "加课码", "VARCHAR(8)", "", "学生加入凭证"),
        ("semester", "学年学期", "VARCHAR(64)", "", ""),
        ("is_pinned", "是否置顶", "TINYINT", "", "0/1"),
        ("school_class_id", "行政班ID", "INT", "外键", ""),
    ]),
    ("account_course", "用户课程关联表", [
        ("account", "账号", "VARCHAR", "主键", ""),
        ("class_id", "课程ID", "BIGINT", "主键", ""),
        ("is_archived", "是否归档", "TINYINT", "", "0/1"),
        ("sort_order", "排序", "INT", "", "拖拽排序"),
    ]),
    ("homework", "作业表", [
        ("homework_id", "作业ID", "BIGINT", "主键", "自增"),
        ("name", "作业名称", "VARCHAR(100)", "", ""),
        ("deadline", "截止时间", "DATETIME", "", ""),
        ("type", "作业类型", "ENUM", "", "个人/团队"),
        ("details", "作业描述", "TEXT", "", ""),
        ("attachment_url", "附件路径", "VARCHAR(512)", "", ""),
    ]),
    ("content", "作业提交表", [
        ("content_id", "作业ID", "BIGINT", "主键", "关联 homework"),
        ("account", "学生账号", "VARCHAR", "主键", ""),
        ("score", "分数", "TINYINT", "", ""),
        ("details", "提交内容", "TEXT", "", ""),
        ("is_graded", "是否批阅", "TINYINT", "", ""),
        ("attachment_url", "附件", "VARCHAR(512)", "", ""),
    ]),
    ("course_activity", "课程活动表", [
        ("id", "活动ID", "BIGINT", "主键", ""),
        ("class_id", "课程ID", "BIGINT", "外键", ""),
        ("type", "类型", "VARCHAR(32)", "", "interaction/topic/test等"),
        ("title", "标题", "VARCHAR(255)", "", ""),
        ("content", "内容", "TEXT", "", ""),
        ("deadline", "截止时间", "DATETIME", "", ""),
        ("publish_status", "发布状态", "VARCHAR(16)", "", "draft/published"),
        ("interaction_kind", "互动类型", "VARCHAR(16)", "", "qa/vote/race"),
        ("interaction_options", "互动状态JSON", "TEXT", "", "status/round等"),
    ]),
    ("course_interaction_response", "互动回答表", [
        ("id", "ID", "BIGINT", "主键", ""),
        ("activity_id", "活动ID", "BIGINT", "外键", ""),
        ("account", "账号", "VARCHAR(64)", "", ""),
        ("content", "回答内容", "TEXT", "", ""),
        ("round_num", "轮次", "INT", "", ""),
        ("option_index", "投票选项", "INT", "", "投票时使用"),
    ]),
    ("course_attendance_session", "签到会话表", [
        ("id", "会话ID", "BIGINT", "主键", ""),
        ("class_id", "课程ID", "BIGINT", "外键", ""),
        ("teacher_account", "教师", "VARCHAR(64)", "", ""),
        ("code", "签到码", "VARCHAR", "", ""),
        ("status", "状态", "VARCHAR", "", "open/closed"),
        ("duration_minutes", "时长", "INT", "", "分钟"),
    ]),
    ("course_attendance_record", "签到记录表", [
        ("id", "ID", "BIGINT", "主键", ""),
        ("session_id", "会话ID", "BIGINT", "外键", ""),
        ("account", "学生账号", "VARCHAR(64)", "", ""),
        ("check_time", "签到时间", "TIMESTAMP", "", ""),
    ]),
    ("course_test_question", "测试题目表", [
        ("id", "题目ID", "BIGINT", "主键", ""),
        ("activity_id", "测试活动ID", "BIGINT", "外键", ""),
        ("question_type", "题型", "VARCHAR(16)", "", "choice/short"),
        ("stem", "题干", "TEXT", "", ""),
        ("stem_image_url", "题干配图", "VARCHAR(512)", "", ""),
        ("score", "分值", "INT", "", ""),
    ]),
    ("notification", "通知表", [
        ("id", "通知ID", "INT", "主键", ""),
        ("account", "接收账号", "VARCHAR(64)", "", ""),
        ("class_id", "课程ID", "INT", "", ""),
        ("type", "类型", "VARCHAR(32)", "", "homework/interaction等"),
        ("message", "消息内容", "VARCHAR(512)", "", ""),
        ("is_read", "是否已读", "TINYINT", "", ""),
    ]),
    ("teacher_prep_item", "备课区表", [
        ("id", "ID", "BIGINT", "主键", ""),
        ("teacher_account", "教师", "VARCHAR(64)", "", ""),
        ("kind", "类型", "VARCHAR(32)", "", "homework/topic/test等"),
        ("title", "标题", "VARCHAR(255)", "", ""),
        ("meta_json", "扩展JSON", "TEXT", "", "题型配置等"),
    ]),
    ("course_grade_weight", "成绩权重表", [
        ("class_id", "课程ID", "BIGINT", "主键", "", ""),
        ("homework_weight", "作业权重", "INT", "", "百分比"),
        ("test_weight", "测试权重", "INT", "", ""),
        ("interaction_weight", "互动权重", "INT", "", ""),
        ("attendance_weight", "考勤权重", "INT", "", ""),
    ]),
]


def build_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("软件工程与计算 II 课程设计报告")
    set_run_font(r, "黑体", 22, True)
    for _ in range(2):
        doc.add_paragraph()
    items = [
        ("题    目", "课堂派（ClassPi）"),
        ("学    院", "两江人工智能学院"),
        ("专    业", "软件工程"),
        ("学生姓名", "（填写）          学号  （填写）"),
        ("学生姓名", "（填写）          学号  （填写）"),
        ("学生姓名", "（填写）          学号  （填写）"),
        ("指导教师", "王森、张金荣、龙华、蒋鑫等"),
        ("时    间", "2026 年 06 月"),
    ]
    for label, val in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = Pt(28)
        r = p.add_run(f"{label}    {val}")
        set_run_font(r, "宋体", 14)


def build_team_tables(doc):
    doc.add_page_break()
    add_section(doc, "团队职责分工表")
    add_table(doc, ["姓名", "角色", "工作职责"], [
        ["（填写）", "项目经理", "开发计划制定与执行、系统设计、课程与互动模块开发、文档撰写"],
        ["（填写）", "前后端开发", "作业管理、在线测试、签到考勤模块开发、接口联调"],
        ["（填写）", "前后端开发", "个人设置、消息通知、备课区、成绩册模块开发、系统测试"],
    ])
    add_section(doc, "团队沟通计划表")
    add_table(doc, ["沟通场景", "沟通方式", "沟通频率", "沟通责任人", "沟通内容"], [
        ["集中上课", "当面沟通", "每次课", "项目经理", "设计评审、实现问题、进度同步"],
        ["非集中上课", "线上/当面", "4 次/天", "项目经理", "Bug 修复、接口联调、代码合并"],
        ["里程碑节点", "组会演示", "每周 1 次", "全体成员", "功能演示、任务分配、风险识别"],
    ])


def build_abstract(doc):
    doc.add_page_break()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(h.add_run("摘  要"), "黑体", 16, True)
    add_body(doc,
        "随着教育信息化的深入发展，传统的教学模式正面临着转型升级的迫切需求。"
        "本文提出并实现了一个基于 Vue 3 前端框架和 Spring Boot 3 后端技术的互动式教学管理系统——课堂派（ClassPi），"
        "旨在通过现代化的技术手段提升教学互动性和管理效率。课堂派系统集成了实时课堂互动（问答、投票、抢答、随机点名）、"
        "在线签到考勤、作业发布与批阅、在线测试测验、教师备课区、消息通知、可配置成绩册等功能模块。"
        "系统采用前后端分离的 B/S 架构，使用 JWT 进行身份认证、BCrypt 加密存储密码、WebSocket 实现课堂实时事件推送，"
        "MyBatis 访问 MySQL 数据库完成数据持久化。在开发过程中，团队采用敏捷迭代方式，"
        "实现了从需求分析、系统设计、编码实现到系统测试的完整软件工程流程。"
        "经过功能测试与联调验证，课堂派运行稳定，有效提升了课堂参与度与教学管理效率，"
        "为高校智慧教学提供了可落地的解决方案。"
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    set_run_font(p.add_run("关键词："), "宋体", 12, True)
    set_run_font(p.add_run("Vue.js；Spring Boot；课堂派；互动式教学；前后端分离；WebSocket"), "宋体", 12)

    doc.add_paragraph()
    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(h2.add_run("Abstract"), "黑体", 16, True)
    add_body(doc,
        "With the deepening development of educational informatization, traditional teaching models "
        "are facing an urgent need for transformation. This paper proposes and implements ClassPi, "
        "an interactive teaching management system based on Vue 3 and Spring Boot 3. The platform "
        "integrates real-time classroom interactions, attendance check-in, homework management, "
        "online tests, teacher preparation area, notifications, and configurable grade books. "
        "Through a front-end and back-end separation architecture with JWT authentication and "
        "WebSocket live events, the system achieves high responsiveness and security. "
        "Practical deployment demonstrates significant improvements in classroom engagement "
        "and teaching management efficiency."
    )
    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = Cm(0.74)
    set_run_font(p2.add_run("Keywords: "), "宋体", 12, True)
    set_run_font(p2.add_run("Vue.js; Spring Boot; ClassPi; Interactive Teaching; WebSocket"), "宋体", 12)


def build_toc(doc):
    doc.add_page_break()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(h.add_run("目  录"), "黑体", 16, True)
    doc.add_paragraph()
    add_toc_field(doc)


def build_chapter1(doc, images):
    add_chapter(doc, "1  需求分析")
    add_section(doc, "1.1  需求概述")
    add_body(doc,
        "课堂派（ClassPi）是一款面向高等院校的在线教学管理平台，适用于面对面课堂与线上混合式教学场景。"
        "平台为教师和学生提供了一个互动性强、功能完善的数字化教学环境。以下从功能与非功能两个维度概述系统需求。"
    )
    add_numbered_items(doc, [
        "用户登录验证：提供稳定的 JWT 认证机制，密码 BCrypt 加密存储，支持教师与学生角色区分及权限控制。",
        "用户信息管理：支持注册、编辑个人资料（姓名、院系、专业、年级、学号/工号、入学日期等）、修改密码与头像上传。",
        "教师创建课程：教师可创建课程，设置课程名称、上课时间、关联行政班、学年学期，系统自动生成加课码。",
        "学生加入课程：学生通过加课码或选课加入课程，查看课程成员与各类教学活动。",
        "课程管理：支持课程置顶、归档、拖拽排序、按学期筛选，便于师生快速定位课程。",
        "作业管理：教师布置作业（含附件），设定截止日期；学生在线提交；教师批阅、打分、打回。",
        "课堂互动：支持实时问答、投票、抢答、随机点名；WebSocket 实时同步；结束签到时自动关闭进行中的互动。",
        "签到考勤：教师发起课堂签到，生成签到码；学生打卡；统计应到、实到、缺勤人数并保留历史记录。",
        "在线测试：支持选择题与简答题，可保存草稿后发布，自动统计交卷人数与答题情况。",
        "课程活动：支持话题讨论、资料发布、公告通知，学生可回复与参与讨论。",
        "教师备课区：教师可预制作业、资料、话题、测试，一键导入到具体课程，提高备课效率。",
        "消息通知：作业发布、话题更新、签到开启、互动发布、测试发布等事件实时推送通知。",
        "成绩册：教师可配置作业、测试、互动、考勤权重，系统自动汇总学生综合成绩。",
    ])

    add_section(doc, "1.2  总体需求描述")
    add_body(doc,
        "课堂派的目标是为教育者和学习者打造一个功能全面、操作便捷的在线教学平台。"
        "教师可通过平台管理课程、布置与批改作业、发起签到与课堂互动、发布测试与资料、使用备课区提高效率；"
        "学生可通过平台加入课程、提交作业、参与互动与讨论、完成签到、查看成绩与通知。"
        "系统应保证长时间稳定运行，界面简洁直观，操作流程清晰，降低用户学习成本，并具备良好的可扩展性。"
    )
    add_body(doc,
        "在总体业务流程上，教师登录后进入课程列表，选择或创建课程后进入课程主页，"
        "可在不同 Tab 页中管理作业、互动、测试、资料、话题、成员与成绩册；"
        "学生登录后查看已加入课程，接收通知并参与各类教学活动。"
        "「进入课堂」页面集中展示当前签到状态与课堂互动入口，实现当堂教学活动的统一管理。"
    )

    add_section(doc, "1.3  详细业务需求")
    add_subsection(doc, "1.3.1  角色权限管理")
    add_body(doc,
        "课堂派通过角色权限管理确保不同用户只能访问其权限范围内的功能。"
        "教师角色拥有创建与管理课程、布置与批改作业、发起签到与互动、发布测试、管理成绩册、使用备课区等权限；"
        "学生角色拥有加入课程、提交作业、参与互动与话题、签到打卡、查看个人成绩与通知等权限。"
        "系统在每次 API 请求时通过 JWT 令牌识别用户身份，并在 Service 层进一步校验操作权限。"
    )
    add_figure(doc, images["role_activity"], "图 1.1  角色权限活动图")
    add_use_case_table(doc, "CP-001", "教师/学生", "角色权限管理",
                       "用户已登录系统",
                       "1.获取角色权限 2.维护个人信息 3.加载对应菜单 4.执行授权操作",
                       "不同角色看到不同功能界面", "表 1-1  角色权限管理用例描述")

    add_subsection(doc, "1.3.2  课程管理")
    add_body(doc,
        "课程管理是系统的核心模块之一。教师创建课程后系统生成唯一加课码，"
        "可将课程关联一个或多个行政班，学生通过加课码或行政班自动关联加入课程。"
        "师生均可对课程进行置顶或归档操作；教师还可通过拖拽调整课程在列表中的显示顺序。"
        "课程主页聚合展示该课程下的作业、互动、测试、资料、话题、成员与成绩册等子模块。"
    )
    add_figure(doc, images["course_activity"], "图 1.2  课程管理活动图")
    add_use_case_table(doc, "CP-002", "教师", "创建课程", "教师已登录",
                       "1.点击创建课程 2.填写课程信息 3.确认创建 4.获得加课码",
                       "课程列表新增该课程", "表 1-2  创建课程用例描述")
    add_use_case_table(doc, "CP-003", "学生", "加入课程", "学生已登录",
                       "1.点击加入课程 2.输入加课码 3.验证通过 4.进入课程主页",
                       "课程列表展示新加入课程", "表 1-3  加入课程用例描述")

    add_subsection(doc, "1.3.3  作业管理")
    add_body(doc,
        "作业管理模块支持教师发布个人作业与团队作业，设定截止日期与作业说明，可附带参考文件。"
        "学生在截止日前提交作业内容与附件，系统记录提交时间并支持更新提交。"
        "教师可在作业详情页查看每位学生的提交状态（已交/未交/已批阅），进行打分、写评语或打回重做。"
        "作业相关的通知会推送给尚未提交的学生，教师也可手动催交。"
    )

    add_subsection(doc, "1.3.4  课堂互动与签到")
    add_body(doc,
        "课堂互动模块支持三种互动形式：实时问答（多轮提问+随机点名）、投票（单选+实时统计）、抢答（教师开启后学生抢先回答）。"
        "互动状态存储于 course_activity.interaction_options JSON 字段，通过 WebSocket 向在线用户推送更新。"
        "签到模块与互动模块联动：教师结束签到时，系统自动关闭该课程所有进行中的课堂互动，"
        "避免当堂互动在课后仍显示「进行中」。"
    )

    add_subsection(doc, "1.3.5  在线测试与备课区")
    add_body(doc,
        "在线测试支持选择题（A/B/C/D 四选一）与简答题，教师可先保存草稿再发布。"
        "学生交卷后教师可查看答卷并批改简答题。备课区允许教师预先创建作业、资料、话题、测试模板，"
        "在课程中通过「从备课区导入」一键发布，支持测试题含简答题与题干配图。"
    )

    add_section(doc, "1.4  非功能性需求")
    nfr_sections = [
        ("1.4.1  环境需求", "开发环境：Java 17、Spring Boot 3.4、MyBatis 3.0、Vue 3、Vite、MySQL 8.0、Maven、IntelliJ IDEA、Node.js。"
         "运行环境：Windows 10 及以上，Chrome/Edge 浏览器，后端端口 9090，前端开发端口 5174。"),
        ("1.4.2  性能需求", "页面响应时间平均不超过 2 秒，WebSocket 消息推送延迟低于 1 秒；"
         "支持单课程 100 名以上学生同时在线签到与互动。"),
        ("1.4.3  安全需求", "用户密码 BCrypt 加密存储；JWT 令牌认证；AuthInterceptor 拦截未授权请求；"
         "教师只能操作自己创建的课程，学生只能提交自己的作业与互动回答。"),
        ("1.4.4  兼容需求", "支持 Windows 操作系统与 Chrome、Edge 等主流浏览器；"
         "前端采用响应式布局，适配常见笔记本屏幕分辨率。"),
        ("1.4.5  交互需求", "统一 AppModal 弹窗组件、Toast 消息提示、Confirm 确认框；"
         "课程列表支持拖拽排序动画；互动与签到状态 WebSocket 实时刷新。"),
        ("1.4.6  扩展需求", "预留 AI 辅助出题与批阅接口（AiController）；"
         "成绩权重可配置；行政班 Excel 批量导入；附件在线预览与下载。"),
    ]
    for title, text in nfr_sections:
        add_subsection(doc, title)
        add_body(doc, text)
    add_body(doc,
        "在安全审计方面，系统通过 Spring Boot 日志记录关键操作异常；数据库迁移脚本（migration_*.sql）"
        "版本化管理表结构变更，便于团队拉代码后执行 bootstrap_after_pull.sql 同步环境。"
        "在灾难恢复方面，MySQL 定期备份可恢复业务数据；uploads 目录需同步备份以保留用户上传的附件。"
        "系统设计与运营符合学校信息化管理要求，注重用户隐私与数据安全。"
    )


def build_chapter2(doc, images):
    add_chapter(doc, "2  系统设计")
    add_section(doc, "2.1  设计概述")
    for para in design_overview_paragraphs():
        add_body(doc, para)

    add_section(doc, "2.2  总体设计")
    add_subsection(doc, "2.2.1  功能架构设计")
    add_body(doc,
        "平台包含用户管理、课程管理、作业管理、课堂互动、签到考勤、在线测试、"
        "课程活动（话题/资料/公告）、教师备课区、消息通知、成绩册十大功能模块，"
        "共计 30 余个功能点，模块间通过统一 REST API 与 WebSocket 通道协作。"
    )
    add_figure(doc, images["func_tree"], "图 2.1  平台功能架构图")
    add_table(doc, ["功能模块", "功能点", "说明"], [
        ["用户管理", "注册/登录/JWT/个人信息", "身份认证与权限控制"],
        ["课程管理", "创建/加入/置顶/归档/排序", "课程全生命周期管理"],
        ["作业管理", "发布/提交/批阅/打回/催交", "含附件上传与预览"],
        ["课堂互动", "问答/投票/抢答/点名", "WebSocket 实时同步"],
        ["签到考勤", "发起/打卡/统计/历史", "应到实到缺勤"],
        ["在线测试", "选择/简答/草稿/批改", "自动统计交卷"],
        ["备课区", "预制/导入/配图", "提高备课效率"],
        ["成绩册", "权重/汇总/导出", "多维度成绩计算"],
    ], "表 2-1  系统功能模块表")

    add_subsection(doc, "2.2.2  技术架构设计")
    add_figure(doc, images["arch"], "图 2.2  系统技术架构图")
    add_body(doc,
        "展示层：Vue 3 单页应用，通过 Vue Router 管理路由，Pinia 管理全局状态。"
        "交互层：Axios 发送 HTTP 请求至 /editor 前缀 API；WebSocket 连接 /ws/live 订阅实时事件。"
        "业务层：Spring Boot Controller 接收请求，Service 处理业务逻辑，JWT 拦截器校验身份。"
        "持久层：MyBatis Mapper 操作 MySQL；FileStorageService 管理 uploads 目录下的附件文件。"
    )

    add_subsection(doc, "2.2.3  物理模型设计")
    add_body(doc,
        "开发阶段：前端 Vite Dev Server（5174）通过 proxy 转发 API 至后端（9090）。"
        "部署阶段：前端构建为静态资源可由 Nginx 托管；Spring Boot 内嵌 Tomcat 运行后端服务；"
        "MySQL 存储业务数据；uploads 目录存储用户上传的头像、作业附件、资料文件等。"
    )

    add_section(doc, "2.3  数据库设计")
    add_subsection(doc, "2.3.1  概念结构设计")
    add_body(doc,
        "根据需求分析建立 E-R 概念模型。主要实体包括：用户（accounts）、课程（courses）、"
        "用户-课程关联（account_course）、作业（homework）、作业提交（content）、"
        "课程活动（course_activity）、互动回答（course_interaction_response）、"
        "签到会话（course_attendance_session）、签到记录（course_attendance_record）、"
        "测试题目（course_test_question）、通知（notification）、备课项（teacher_prep_item）、"
        "成绩权重（course_grade_weight）等。实体间通过外键与关联表建立联系。"
    )
    add_figure(doc, images["er"], "图 2.3  课堂派核心 E-R 图")
    add_body(doc,
        "下图所示为 ClassPi 后端 Controller 层接口划分（节选），以及前端 WebSocket 连接实现，"
        "均来自本项目源代码截图，而非第三方示例。"
    )
    if images.get("code_course"):
        add_figure(doc, images["code_course"], "图 2.4  CourseController 接口代码（ClassPi 项目）")
    if images.get("code_websocket"):
        add_figure(doc, images["code_websocket"], "图 2.5  WebSocket 连接代码（ClassPi 项目）")

    add_subsection(doc, "2.3.2  数据库逻辑结构设计")
    add_body(doc,
        "将 E-R 图转换为关系模型，遵循第三范式减少数据冗余。"
        "用户与课程为多对多关系，通过 account_course 关联表实现；"
        "课程与作业为一对多关系，通过 courses_homework 关联；"
        "课程活动统一存储于 course_activity 表，通过 type 字段区分互动/话题/测试/资料/公告。"
    )

    add_subsection(doc, "2.3.3  数据库物理结构设计")
    add_body(doc,
        "选用 MySQL 8.0 作为数据库管理系统，字符集 utf8mb4，存储引擎 InnoDB。"
        "对高频查询字段建立索引，如 course_activity(class_id, type)、"
        "course_interaction_response(activity_id, round_num) 等。"
        "时间字段统一使用 TIMESTAMP/DATETIME，由应用层 HomeworkDeadlineUtil 格式化显示。"
    )

    add_subsection(doc, "2.3.4  数据表结构描述")
    for idx, (table, cn, fields) in enumerate(DB_TABLES, 1):
        rows = [(f[0], f[1], f[2], f[3], f[4]) for f in fields]
        add_table(doc, ["字段", "中文名", "类型", "键", "备注"], rows,
                  caption=f"表 2-{idx}  {cn}（{table}）")

    add_section(doc, "2.4  详细设计")
    add_subsection(doc, "2.4.1  用户管理模块")
    add_lines(doc, module_design_user())
    add_figure(doc, images["seq_login"], "图 2.34  用户登录时序图（ClassPi 实现）")

    add_subsection(doc, "2.4.2  课程管理模块")
    add_lines(doc, module_design_course())

    add_subsection(doc, "2.4.3  作业管理模块")
    add_lines(doc, module_design_homework())
    add_figure(doc, images["seq_homework"], "图 2.35  作业提交时序图（ClassPi 实现）")

    add_subsection(doc, "2.4.4  课堂互动模块")
    add_lines(doc, module_design_interaction())
    add_figure(doc, images["seq_interaction"], "图 2.36  课堂互动时序图（ClassPi 实现）")

    add_subsection(doc, "2.4.5  签到考勤模块")
    add_body(doc,
        "AttendanceController 提供 startAttendance、checkInAttendance、closeAttendance、"
        "getOpenAttendance、listAttendances 等接口。发起签到时若已有 open 会话则先关闭再新建，"
        "并同步关闭进行中的课堂互动。生成 4 位签到码，计算应到人数（关联行政班学生 + 已加课学生）。"
        "学生打卡写入 course_attendance_record，教师可查看实到/缺勤名单。关闭签到时 publish attendance_closed 事件，"
        "前端 LiveClass.vue 与 InteractionContent.vue 订阅 course 房间并刷新状态。"
    )

    add_subsection(doc, "2.4.6  在线测试与备课区模块")
    add_body(doc,
        "TestController 支持 addCourseTest、saveCourseTestDraft、submitTest、gradeTestAnswer、deleteTest。"
        "测试题目存于 course_test_question，答卷存于 course_test_submission 与 course_test_answer。"
        "PrepController 支持 savePrepItem、listPrepItems、publishPrepToCourse，"
        "可将备课区测试（含简答题与 stem_image_url 题干配图）导入为课程测试活动。"
        "PrepImportModal 在 courseContent.vue 中触发，使用 elevated 属性确保弹窗层级高于其他 Modal。"
    )

    add_subsection(doc, "2.4.7  话题与资料模块")
    add_body(doc,
        "ActivityController 统一管理 type 为 topic、material、announcement 的课程活动。"
        "教师发布话题/资料/公告，学生可查看与回复。course_activity_reply 存储讨论回复，支持附件。"
        "NotificationService 在发布作业、话题、签到、测试、互动等事件时写入 notification 表，"
        "前端 mainClass.vue 轮询或进入页面时拉取未读通知并弹窗提示。"
    )


def build_chapter3(doc, images):
    add_chapter(doc, "3  系统实现")
    add_section(doc, "3.1  系统实现概述")
    for para in impl_overview_paragraphs():
        add_body(doc, para)

    add_section(doc, "3.2  代码结构")
    add_figure(doc, images["frontend"], "图 3.1  前端项目代码结构图")
    add_figure(doc, images["backend"], "图 3.2  后端项目代码结构图")
    add_body(doc,
        "后端采用经典三层架构：Controller 层接收 HTTP 请求并返回 JSON；"
        "Service 层封装业务逻辑与事务；Mapper 层通过 MyBatis 注解或 XML 操作数据库。"
        "dto 包封装请求与响应对象，entity 包对应数据库表实体，util 包提供文件存储、"
        "截止日期格式化等工具。security 包含 JwtService 与 AuthInterceptor。"
        "live 包包含 LiveEventHub 与 LiveEventPublisher 实现 WebSocket 广播。"
    )

    add_section(doc, "3.3  关键功能实现")
    add_subsection(doc, "3.3.1  文件上传和预览")
    add_body(doc,
        "平台实现了统一的文件上传机制。前端通过 FormData 发送 multipart 请求，"
        "后端 FileStorageService 将文件保存至 uploads 目录并返回访问 URL。"
        "作业附件、资料附件、头像、测试题干配图等均复用该机制。"
        "预览时前端通过 window.open(url) 或 iframe 打开文件 URL。"
    )

    add_subsection(doc, "3.3.2  用户注册登录")
    add_body(doc,
        "注册时前端提交账号、密码、姓名、角色、机构等信息，后端验证账号唯一性后 BCrypt 编码密码入库。"
        "登录成功后返回 JWT Token，前端存储于 Pinia 并附加到后续请求的 Authorization 头。"
        "公开页面（登录/注册）不携带 Token，避免登录过期时无法访问注册页的问题。"
    )
    add_figure(doc, images["code_jwt"], "图 3.3  JWT Token 生成关键代码（ClassPi 项目源码）")

    add_subsection(doc, "3.3.3  作业管理功能")
    add_body(doc,
        "作业发布：教师填写作业名称、类型、截止日期、描述，可上传附件，后端写入 homework 与 courses_homework。"
        "作业提交：学生在 HomeworkContent.vue 填写内容并上传附件，写入 content 表。"
        "作业批阅：教师查看提交列表，进入详情页打分写评语，更新 is_graded 标记。"
        "作业删除：Service 层事务内删除 homework、courses_homework、content 关联数据。"
    )

    add_subsection(doc, "3.3.4  课堂互动与 WebSocket")
    add_body(doc,
        "前端 InteractionContent.vue 通过 connectLiveSocket 订阅 interaction:{activityId} 与 course:{classId} 房间。"
        "收到 interaction_updated、interaction_closed、attendance_closed 等事件后自动刷新详情。"
        "教师可发布多轮问答、开启抢答、随机点名；学生提交回答或投票；"
        "结束签到时后端批量关闭 active 状态互动并广播 interactions_closed 事件。"
    )
    add_figure(doc, images["code_close"], "图 3.4  结束签到自动关闭互动关键代码（ClassPi 项目源码）")
    if images.get("code_interaction"):
        add_figure(doc, images["code_interaction"], "图 3.5  课堂互动发布关键代码（ClassPi 项目源码）")
    if images.get("code_interaction_vue"):
        add_figure(doc, images["code_interaction_vue"], "图 3.6  互动页 WebSocket 订阅代码（ClassPi 项目源码）")

    add_subsection(doc, "3.3.5  备课区与测试")
    add_body(doc,
        "PrepArea.vue 提供备课项 CRUD，测试类备课项支持选择题与简答题及题干图片上传。"
        "PrepImportModal 在课程主页触发，调用 publishPrepToCourse 将备课项复制为课程活动。"
        "TestEditor.vue 支持编辑题目、设置正确答案与分值，可保存草稿后发布。"
    )

    add_section(doc, "3.4  系统实现界面展示")
    add_body(doc,
        "以下界面截图均在 ClassPi 系统本地运行环境（前端 localhost:5174、后端 localhost:9090）"
        "中使用 Playwright 自动截取，展示的是本项目真实页面，而非参考文档或占位图。"
    )
    add_subsection(doc, "3.4.1  角色登录注册")
    add_body(doc, "用户通过登录页输入账号密码，系统识别教师/学生身份并跳转对应主页。无账号时在注册页选择角色完成注册。")
    add_figure(doc, images["ui_login"], "图 3.7  角色登录界面（ClassPi 实际截图）")
    add_figure(doc, images["ui_register"], "图 3.8  角色注册界面（ClassPi 实际截图）")

    add_subsection(doc, "3.4.2  主页面及功能")
    add_body(doc, "教师主页可创建课程与加入他人课程；学生主页仅可加入课程。支持学期筛选、置顶、归档与拖拽排序。")
    add_figure(doc, images["ui_main_teacher"], "图 3.9  教师主界面（ClassPi 实际截图）")
    add_figure(doc, images["ui_main_student"], "图 3.10  学生主界面（ClassPi 实际截图）")

    add_subsection(doc, "3.4.3  课程主页与进入课堂")
    add_body(doc, "课程主页按 Tab 展示互动、话题、资料、测试、公告、作业、成员、成绩册。签到进行中时顶部显示横幅提示。")
    add_figure(doc, images["ui_course"], "图 3.11  教师课程主页（ClassPi 实际截图）")
    if images.get("ui_course_student"):
        add_figure(doc, images["ui_course_student"], "图 3.12  学生课程主页（ClassPi 实际截图）")
    add_figure(doc, images["ui_live"], "图 3.13  进入课堂/签到界面（ClassPi 实际截图）")

    add_subsection(doc, "3.4.4  课堂互动与作业")
    if images.get("ui_interaction"):
        add_figure(doc, images["ui_interaction"], "图 3.14  课堂互动界面（ClassPi 实际截图）")
    add_figure(doc, images["ui_homework"], "图 3.15  课程作业列表（ClassPi 实际截图）")

    add_subsection(doc, "3.4.5  备课区、成绩册与设置")
    add_figure(doc, images["ui_prep"], "图 3.16  教师备课区界面（ClassPi 实际截图）")
    if images.get("ui_grade"):
        add_figure(doc, images["ui_grade"], "图 3.17  成绩册界面（ClassPi 实际截图）")
    add_figure(doc, images["ui_setting"], "图 3.18  个人设置界面（ClassPi 实际截图）")


def build_chapter4(doc):
    add_chapter(doc, "4  系统测试")
    add_section(doc, "4.1  测试概述")
    for para in test_overview_paragraphs():
        add_body(doc, para)
    add_section(doc, "4.2  测试环境")
    add_table(doc, ["项目", "配置"], [
        ["操作系统", "Windows 11"],
        ["浏览器", "Google Chrome 120+"],
        ["数据库", "MySQL 8.0，库名 t_class"],
        ["后端", "Spring Boot 3.4，端口 9090"],
        ["前端", "Vue 3 + Vite，端口 5174"],
    ], "表 4-1  测试环境配置表")

    add_section(doc, "4.3  测试方法")
    add_body(doc,
        "采用黑盒功能测试为主，按测试用例步骤执行操作并记录实际结果。"
        "对 WebSocket 实时功能进行多浏览器窗口并行测试。"
        "对边界条件（空输入、过期 Token、重复提交、截止后提交）进行异常测试。"
    )

    add_section(doc, "4.4  测试用例设计")
    cases = [
        ("CP-T01", "账号注册与登录", "1.点击注册 2.填写信息 3.提交 4.登录",
         "账号、密码、姓名、角色", "注册成功并可登录跳转主页"),
        ("CP-T02", "创建与加入课程", "1.教师创建课程 2.获取加课码 3.学生加入",
         "课程名称、班级、学期", "双方课程列表均显示该课程"),
        ("CP-T03", "发布与提交作业", "1.教师添加作业 2.学生提交 3.查看状态",
         "作业名称、截止日期、附件", "作业列表显示，提交成功"),
        ("CP-T04", "批阅作业", "1.教师打开提交 2.打分写评语 3.保存",
         "分数、评语", "学生可查看成绩"),
        ("CP-T05", "课堂互动问答", "1.教师开启互动 2.发布问题 3.学生回答",
         "问题文本、回答文本", "实时显示回答列表"),
        ("CP-T06", "签到考勤", "1.教师发起签到 2.学生打卡 3.查看统计",
         "签到码", "应到实到数据正确"),
        ("CP-T07", "结束签到关闭互动", "1.开启互动 2.发起签到 3.结束签到",
         "无", "互动状态变为已结束"),
        ("CP-T08", "在线测试", "1.创建测试 2.学生交卷 3.查看统计",
         "选择题/简答题答案", "交卷人数正确"),
        ("CP-T09", "备课区导入", "1.创建备课项 2.导入课程 3.验证",
         "备课测试含简答题", "课程中出现对应活动"),
        ("CP-T10", "附件上传下载", "1.上传附件 2.预览 3.下载",
         "doc/pdf 文件", "文件内容一致无乱码"),
    ]
    for i, (num, name, steps, inp, exp) in enumerate(cases, 1):
        add_test_case_table(doc, num, name, steps, inp, exp, f"表 4-{i+1}  {name}测试用例")


def build_chapter5(doc):
    add_chapter(doc, "5  总结")
    for para in summary_paragraphs():
        add_body(doc, para)

    add_section(doc, "参考文献")
    refs = [
        "[1] 李鸿君. 大话软件工程——需求分析与软件设计[M]. 北京: 清华大学出版社, 2020.",
        "[2] 刘慧娟. 基于 SpringBoot 的民主测评系统的设计与实现[D]. 北京邮电大学, 2022.",
        "[3] 胡强. MySQL 数据库常见问题分析与研究[J]. 电脑编程技巧与维护, 2019.",
        "[4] 郑海燕. 基于 Java Web 的高校英语线上教学平台设计[J]. 自动化与仪器仪表, 2023.",
        "[5] 时俊雅, 黄苏雨. 基于 Java 的大学生在线学习系统设计与实现[J]. 无线互联科技, 2023.",
        "[6] 孙洪盼. 基于 SpringBoot 和 Vue 的友为交流社区的设计与实现[D]. 重庆大学, 2022.",
        "[7] 韩坤, 林关成, 安嘉豪等. 基于 B/S 的高校学生学业预警系统设计与实现[J]. 信息技术, 2021.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        set_run_font(p.add_run(ref), "宋体", 12)


def main():
    import os
    if os.environ.get("SKIP_CAPTURE") != "1":
        print("Capturing real project screenshots...")
        capture_project_assets()
    else:
        print("Skipping capture (SKIP_CAPTURE=1)")
    print("Loading images...")
    images = load_project_images()

    missing = [k for k in ("ui_login", "ui_main_teacher", "ui_course") if k not in images or not Path(images[k]).exists()]
    if missing:
        print(f"WARNING: missing screenshots {missing}. Ensure frontend:5174 and backend:9090 are running.")

    doc = Document()
    setup_document_styles(doc)

    # Section 1: cover — no header/footer
    setup_page(doc.sections[0], header_text=None, footer_page_num=False)
    build_cover(doc)
    build_team_tables(doc)

    # Section 2: abstract
    doc.add_section(WD_SECTION.NEW_PAGE)
    setup_page(doc.sections[-1], header_text=HEADER, footer_page_num=True)
    build_abstract(doc)

    # Section 3: TOC
    doc.add_section(WD_SECTION.NEW_PAGE)
    setup_page(doc.sections[-1], header_text=HEADER, footer_page_num=True)
    build_toc(doc)

    # Section 4: main body
    doc.add_section(WD_SECTION.NEW_PAGE)
    setup_page(doc.sections[-1], header_text=HEADER, footer_page_num=True)
    build_chapter1(doc, images)
    build_chapter2(doc, images)
    build_chapter3(doc, images)
    build_chapter4(doc)
    build_chapter5(doc)

    doc.save(str(OUT))
    try:
        doc.save(str(OUT_ALT))
        print(f"Saved: {OUT_ALT}")
    except PermissionError:
        print(f"Skip {OUT_ALT} (file open in Word)")
    try:
        doc.save(str(OUT_CN))
        print(f"Saved: {OUT_CN}")
    except PermissionError:
        print(f"Skip {OUT_CN} (file open in Word)")
    print(f"Saved: {OUT}")
    print("Tip: Open in Word, right-click TOC -> Update Field (F9) for page numbers.")


if __name__ == "__main__":
    main()
